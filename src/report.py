from docx import Document
from docx.shared import Inches
import os

doc = Document()
doc.add_heading("CPU Usage Prediction — Q3 Report", 0)

doc.add_heading("Workflow", level=1)
doc.add_paragraph("...")  # paste the workflow content

doc.add_heading("Metrics", level=1)
if os.path.exists("metrics/metrics.json"):
    import json
    metrics = json.load(open("metrics/metrics.json"))
    p = doc.add_paragraph()
    for k,v in metrics.items():
        p.add_run(f"{k}: {v}\n")

doc.add_heading("Figures", level=1)
for fname in ["plots/feature_importance.png", "plots/pred_vs_actual.png", "plots/residuals.png", "plots/error_distribution.png"]:
    if os.path.exists(fname):
        doc.add_picture(fname, width=Inches(5))
        doc.add_paragraph(fname)

doc.save("Q3_CPU_Usage_Report.docx")
print("Saved Q3_CPU_Usage_Report.docx")
